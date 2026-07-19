from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


IN_DOC = Path(r"C:/Users/hasee/Desktop/paper_v2.docx")
OUT_DOC = Path(r"C:/Users/hasee/Desktop/paper_v2_revised_complete_split_cf_figures_params_source_fixed.docx")

FIGURES = {
    "module": Path(r"D:/pycharm/rgcl_llm/src/paper_figures/module_ablation_results.png"),
    "aux": Path(r"D:/pycharm/rgcl_llm/src/paper_figures/auxiliary_text_comparison.png"),
    "cf": Path(r"D:/pycharm/rgcl_llm/src/cf_evidence_behavior/cf_toxic_only_paper_academic.png"),
    "cf_dist": Path(r"D:/pycharm/rgcl_llm/src/cf_evidence_behavior/cf_drop_distribution_split.png"),
    "cf_ind": Path(r"D:/pycharm/rgcl_llm/src/cf_evidence_behavior/cf_response_indicators_split.png"),
}


REPLACEMENTS = {
    "3.4 反事实证据学习": "3.3 反事实证据学习",
    "下面分别介绍语义解释生成与编码、解释感知投影与融合、检索引导对比学习以及反事实证据学习。":
        "下面分别介绍语义解释生成与编码、解释感知投影与融合，以及反事实证据学习。其中，检索引导对比学习沿用 RGCL 的基本范式，并在语义增强 embedding 空间中完成近邻检索与对比优化，因此与解释感知融合过程一并说明。",
    "其中，M(·) 表示语义引导的跨模态融合过程。具体而言，该过程首先通过线性层将图像特征、OCR 文本特征和筛选后的解释特征分别展开为若干 token，使原本的全局向量能够以 token 序列形式参与注意力计算。为显式建模图文之间的关系，模型进一步构造两类关系 token：图文差异 token |v~_i - t~_i| 用于刻画图像与文本之间的不一致性，图文乘积 token v~_i ⊙ t~_i 用于刻画二者在特征维度上的对齐关系。随后，图像 token、文本 token、解释 token 和关系 token 被拼接为统一序列，并通过多头注意力进行跨模态交互，使视觉信息能够关注文本、解释和关系线索，文本与解释信息也能够根据视觉内容进行更新。在注意力交互之后，模型进一步以文本 token 和解释 token 作为语义引导信号，对整体 token 序列进行一次语义校准，使融合表示更关注与图文组合含义相关的特征。最后，模型分别对图像、文本、解释和关系 token 进行平均池化，并将池化结果拼接后输入线性投影层，得到语义引导的跨模态交互表示 z_i。lambda_m 为可学习的残差系数。该步骤使解释信息参与图文交互建模，同时通过残差形式保留基础图文表示，避免显式解释直接替代原始图文证据。":
        "其中，M(·) 表示语义引导的跨模态融合过程。需要说明的是，本文所称 token 并非原始图像 patch 或文本词级 token，而是由池化后的全局特征经线性投影得到的一组可学习特征 token，并加入位置参数以形成紧凑的注意力交互序列。具体而言，模型首先将图像特征、OCR 文本特征和筛选后的解释特征分别展开为若干特征 token，使全局向量能够在更细粒度的子空间中参与多头注意力计算。为显式刻画图文关系，模型进一步构造图文差异 token |v~_i - t~_i| 和图文乘积 token v~_i ⊙ t~_i：前者用于表示图像与文本之间的潜在不一致性，后者用于表示二者在特征维度上的对齐关系。随后，图像 token、文本 token、解释 token 和关系 token 被拼接为统一序列，并通过多头注意力进行跨模态交互，使视觉、文本、解释和关系线索能够相互更新。在注意力交互之后，模型以文本 token 和解释 token 作为语义引导信号，对整体序列进行一次语义校准，使融合表示更集中于图文组合含义相关的特征。最后，模型分别对图像、文本、解释和关系 token 进行平均池化，并将池化结果拼接后输入线性投影层，得到语义引导的跨模态交互表示 z_i。lambda_m 为可学习的残差系数。该步骤的作用是让解释信息参与图文关系建模，同时通过残差形式保留基础图文表示，避免显式解释直接替代原始图文证据。",
    "最后，模型将基础表示、FiLM 调制表示和语义残差增强表示进行归一化层级加权组合：":
        "最后，模型将基础表示、FiLM 调制表示和语义残差增强表示进行残差归一化加权组合：",
    "其中，q 为可学习的层级权重。该组合方式使模型能够在不同样本上自适应地平衡基础图文证据、解释调制后的表示和语义残差增强表示。最终表示 x_i 输入残差分类头 f_m，得到用于分类和检索对比学习的 embedding：":
        "其中，q 为可学习的层级权重。该残差归一化组合方式使模型能够在不同样本上自适应地平衡基础图文证据、解释调制后的表示和语义残差增强表示，并避免单一路径在训练早期过度主导最终表示。最终表示 x_i 输入残差分类头 f_m，得到用于分类和检索对比学习的 embedding：",
    "实验在同一组 370 个有害样本上比较完整模型 full 与移除 CF 模块的 no_cf。Mean Drop 和 Median Drop 分别表示置信度下降量的均值与中位数；Drop Rate 表示置信度出现下降的样本比例；Toxic→Non-toxic Flip 表示干预后预测由有害类别翻转为无害类别的样本比例。结果如表 4-5 所示。":
        "实验在同一组 370 个有害样本上比较加入 CF 模块的 with_cf 与移除 CF 模块的 no_cf。Mean Drop 和 Median Drop 分别表示置信度下降量的均值与中位数；Drop Rate 表示置信度出现下降的样本比例；Toxic→Non-toxic Flip 表示干预后预测由有害类别翻转为无害类别的样本比例。结果如表 4-5 所示。",
    "作者待补充：为保证实验可复现，最终版本仍需补充图像编码器、OCR 文本编码器、解释编码器、多模态大模型及版本、优化器、学习率、batch size、温度系数、损失权重、top-k、削弱比例、margin、随机种子、检索库更新频率和硬件环境等细节。":
        "在实现细节上，本文使用预先抽取的图像、OCR 文本和语义解释特征作为模型输入，并将三路特征统一投影到 1024 维隐空间。根据实验脚本设置，训练阶段采用 AdamW 优化器，学习率设为 1e-4，batch size 设为 64，最大训练轮数为 30，梯度裁剪阈值沿用代码默认值 0.1，并固定随机种子为 0。模型采用余弦相似度作为检索和对比学习的距离度量，损失函数采用 triplet-style 检索对比损失，并与二分类交叉熵损失联合优化，其中交叉熵权重为 0.5，triplet margin 为 0.1。训练阶段每个样本检索 1 个 pseudo-gold positive 和 1 个 hard negative，并结合 in-batch negative 构造对比学习目标；检索评估阶段的 top-k 设置为 20，majority voting 方式为 arithmetic。反事实证据学习采用模型实现中的默认超参数：高贡献维度选取比例为 0.12，选中维度的保留系数为 0.05；有害样本置信度下降约束 margin 为 0.23，无害样本约束 margin 为 0.05，辅助反事实损失权重为 0.01。语义引导的跨模态交互模块中，每一路全局特征被展开为 4 个可学习特征 token，并使用 4 头注意力进行交互。",
    "考虑到不同模型结构在优化难度和收敛速度方面存在差异，其达到最优性能的训练轮次可能并不一致。因此，本文将所有模型的最大训练轮数统一设置为 50 个 epoch，以保证各模型获得充分且一致的训练预算。训练过程中，在每个 epoch 结束后于验证集上评估模型性能，并选取验证集表现最佳的 checkpoint 用于测试集评估，从而降低不同收敛进程对模型比较的影响。结果表中的 Best Epoch 表示验证集最优 checkpoint 所对应的训练轮次。":
        "考虑到不同模型结构在优化难度和收敛速度方面存在差异，其达到最优性能的训练轮次可能并不一致。因此，本文按照实验脚本将所有模型的最大训练轮数统一设置为 30 个 epoch，以保证各模型具有一致的训练预算。训练过程中，在每个 epoch 结束后于验证集上评估模型性能，并选取验证集表现最佳的 checkpoint 用于测试集评估，从而降低不同收敛进程对模型比较的影响。结果表中的 Best Epoch 表示验证集最优 checkpoint 所对应的训练轮次。",
    "在关键维度被削弱后，full 的 Mean Drop 为 0.149，明显高于 no_cf 的 0.009；其 Median Drop 也由 0.067 提高至 0.109。这说明加入 CF 模块后，模型所识别的高贡献维度与有害类别预测之间具有更强的对应关系。":
        "在关键维度被削弱后，with_cf 的 Mean Drop 为 0.149，明显高于 no_cf 的 0.009；Median Drop 也由 no_cf 的 0.067 提高至 with_cf 的 0.109。这说明加入 CF 模块后，模型所识别的高贡献维度与有害类别预测之间具有更强的对应关系。",
    "full 的 Drop Rate 为 0.941，即 94.1% 的有害样本在干预后出现置信度下降，而 no_cf 的该比例为 61.6%。此外，full 的 Toxic→Non-toxic Flip 为 31.4%，no_cf 仅为 0.3%。由于上述指标均在定向削弱高贡献维度后计算，较高的置信度下降和类别翻转并不表示模型在正常输入下性能下降，而是表明完整模型对关键证据维度具有更强、更一致的响应。":
        "with_cf 的 Drop Rate 为 0.941，即 94.1% 的有害样本在干预后出现置信度下降，而 no_cf 的该比例为 61.6%。此外，with_cf 的 Toxic→Non-toxic Flip 为 31.4%，no_cf 仅为 0.3%。由于上述指标均在定向削弱高贡献维度后计算，较高的置信度下降和类别翻转并不表示模型在正常输入下性能下降，而是表明完整模型对关键证据维度具有更强、更一致的响应。",
}


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=None):
    run.font.name = west
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), west)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), west)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def reset_para_text(paragraph, text):
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    run = paragraph.add_run(text)
    return run


def paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def is_equation(text):
    stripped = text.strip()
    if not stripped or len(stripped) > 150:
        return False
    if stripped.startswith(("图 ", "表 ", "作者待补充", "Generate ")):
        return False
    markers = ["=", "Σ", "∪", "∈", "≠", "⊙", "softmax", "sigmoid", "Dropout", "SiLU", "LN(", "tanh", "log", "exp("]
    return any(m in stripped for m in markers) and not stripped.endswith("。")


def style_paragraph(paragraph):
    text = paragraph.text.strip()
    name = paragraph.style.name if paragraph.style else ""

    pf = paragraph.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    if name in {"Title_document", "Title"} or paragraph._p.get_or_add_pPr().find(qn("w:pStyle")) is not None and text == "面向中文有害 Meme 检测的解释增强反事实检索对比学习":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, east="黑体", west="Times New Roman", size=18, bold=True)
        return

    if name in {"Subtitle"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, east="Times New Roman", west="Times New Roman", size=12)
        return

    if name.startswith("Heading") or name in {"Head1", "Head2", "AbsHead", "AckHead", "ReferenceHead"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)
        for run in paragraph.runs:
            set_run_font(run, east="黑体", west="Times New Roman", size=13 if name in {"Head1", "Head2"} else 12, bold=True)
        return

    if name == "Caption" or text.startswith(("图 ", "表 ")):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        for run in paragraph.runs:
            set_run_font(run, east="宋体", west="Times New Roman", size=10.5)
        return

    if is_equation(text):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = None
        for run in paragraph.runs:
            set_run_font(run, east="Cambria Math", west="Cambria Math", size=10.5)
            run.italic = True
        return

    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        set_run_font(run, east="宋体", west="Times New Roman", size=10.5)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, east="宋体", west="Times New Roman", size=10, bold=(row_idx == 0))


def add_figure_after(paragraph, image_path, caption, width_cm=14.8):
    img_p = paragraph_after(paragraph)
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(8)
    img_p.paragraph_format.space_after = Pt(3)
    img_p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    cap_p = paragraph_after(img_p, caption, style="Caption")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap_p.runs:
        set_run_font(run, east="宋体", west="Times New Roman", size=10.5)
    return cap_p


def find_para(doc, startswith=None, contains=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if startswith and t.startswith(startswith):
            return p
        if contains and contains in t:
            return p
    raise RuntimeError(f"paragraph not found: {startswith or contains}")


def main():
    doc = Document(str(IN_DOC))

    for p in doc.paragraphs:
        text = p.text.strip()
        if text in REPLACEMENTS:
            reset_para_text(p, REPLACEMENTS[text])

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() == "full":
                    cell.text = "with_cf"

    # Add result figures after the corresponding interpretive paragraphs.
    p_ablation = find_para(doc, contains="二者结合能够带来更稳定的性能收益")
    add_figure_after(
        p_ablation,
        FIGURES["module"],
        "图 4-1 模块消融实验结果。柱状图展示了 Baseline、Baseline + Exp、Baseline + CF 与 SEC-RGCL 在 Acc 和 AUC 上的差异，用于说明语义解释增强与反事实证据学习的独立贡献和联合效果。",
        width_cm=14.6,
    )

    p_aux = find_para(doc, contains="解释文本本身的人类可读质量")
    add_figure_after(
        p_aux,
        FIGURES["aux"],
        "图 4-2 不同辅助文本输入的性能比较。与重复输入原始 OCR 文本相比，语义解释能够提供更有效的辅助语义信息，从而在 Acc 和 AUC 上取得更高结果。",
        width_cm=14.6,
    )

    p_cf = find_para(doc, contains="更强、更一致的响应")
    cf_dist_cap = add_figure_after(
        p_cf,
        FIGURES["cf_dist"],
        "图 4-3 有害样本在反事实证据削弱下的置信度下降分布。横轴表示原始表示与反事实表示之间的有害类别置信度差值，竖线表示零变化位置。",
        width_cm=15.6,
    )
    add_figure_after(
        cf_dist_cap,
        FIGURES["cf_ind"],
        "图 4-4 反事实证据削弱下的汇总响应指标。Mean Drop、Drop Rate 和 Toxic→Non-toxic Flip 分别从平均置信度变化、下降样本比例和类别翻转比例刻画 CF 模块的约束效果。",
        width_cm=15.6,
    )

    for p in doc.paragraphs:
        style_paragraph(p)

    for table in doc.tables:
        style_table(table)

    # Remove stale placeholder styling by making it visually marked but less intrusive.
    for p in doc.paragraphs:
        if p.text.strip().startswith("作者待补充："):
            for run in p.runs:
                set_run_font(run, east="宋体", west="Times New Roman", size=10)
                run.font.color.rgb = RGBColor(128, 0, 0)

    doc.save(str(OUT_DOC))
    check = Document(str(OUT_DOC))
    all_text = "\n".join(p.text for p in check.paragraphs)
    print(OUT_DOC)
    print("paragraphs", len(check.paragraphs), "tables", len(check.tables), "images", len(check.inline_shapes))
    print("contains full:", bool(re.search(r"\\bfull\\b", all_text)))
    print("contains 3.4 heading:", "3.4 反事实证据学习" in all_text)


if __name__ == "__main__":
    main()
