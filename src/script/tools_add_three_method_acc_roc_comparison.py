from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


IN_DOC = Path(r"C:/Users/hasee/Desktop/paper_v2_clean_no_wrong_related_comparison.docx")
OUT_DOC = Path(r"C:/Users/hasee/Desktop/paper_v2_final_three_method_acc_roc_comparison.docx")


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=None):
    run.font.name = west
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), east)
    rpr.rFonts.set(qn("w:ascii"), west)
    rpr.rFonts.set(qn("w:hAnsi"), west)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    return p


def style_body(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run)


def style_heading(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, east="黑体", size=12, bold=True)


def style_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=10.5)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=10, bold=(row_idx == 0))


def find_anchor(doc):
    for p in doc.paragraphs:
        if "由于 AUC 衡量不同阈值下的整体排序能力" in p.text:
            return p
    raise RuntimeError("Could not find section 4.4 anchor.")


def main():
    doc = Document(str(IN_DOC))
    anchor = find_anchor(doc)

    heading = paragraph_after(anchor, "4.4.1 与相关方法的性能比较")
    style_heading(heading)

    intro = paragraph_after(
        heading,
        "为进一步评价本文方法在中文有害 Meme 检测任务中的性能表现，本文选取参考文献 [2] 中具有代表性的多模态方法进行横向比较，包括 Qwen2.5-VL（zero-shot）、Debate-based model 和 FG-E2HMD。"
        "其中，Qwen2.5-VL 代表未针对任务微调的通用多模态大模型，Debate-based model 代表基于解释辩论的多模态推理方法，FG-E2HMD 代表引入细粒度解释增强的中文有害 Meme 检测方法。"
    )
    style_body(intro)

    caption = paragraph_after(intro, "表 4-2a 与相关多模态方法的性能比较")
    style_caption(caption)

    table = doc.add_table(rows=1, cols=4)
    table._tbl.getparent().remove(table._tbl)
    caption._p.addnext(table._tbl)

    headers = ["模型", "类别", "Acc", "ROC"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    rows = [
        ["Qwen2.5-VL（zero-shot）", "Multimodal", "70.25", "73.36"],
        ["Debate-based model", "Multimodal", "79.93", "81.16"],
        ["FG-E2HMD", "Multimodal", "83.53", "83.26"],
        ["SEC-RGCL", "Multimodal", "83.92", "86.22"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    analysis = paragraph_after(caption)
    table._tbl.addnext(analysis._p)
    analysis.add_run(
        "由表 4-2a 可见，SEC-RGCL 在 Acc 和 ROC 两项指标上均取得最高结果。"
        "与 Qwen2.5-VL（zero-shot）相比，SEC-RGCL 的 Acc 和 ROC 分别提高 13.67 和 12.86 个百分点，说明仅依赖通用多模态大模型的零样本推理难以充分适应中文有害 Meme 中的隐含语义。"
        "与 Debate-based model 相比，SEC-RGCL 分别提高 3.99 和 5.06 个百分点，表明将语义解释融入检索引导对比学习能够更有效地优化判别表示。"
        "与 FG-E2HMD 相比，SEC-RGCL 的 Acc 提高 0.39 个百分点，ROC 提高 2.96 个百分点，说明本文在引入语义解释的基础上进一步结合反事实证据学习，有助于提升模型的整体排序判别能力。"
    )
    style_body(analysis)

    doc.save(str(OUT_DOC))
    check = Document(str(OUT_DOC))
    print(OUT_DOC)
    print("paragraphs", len(check.paragraphs), "tables", len(check.tables), "images", len(check.inline_shapes))
    print("has three-method comparison", any("与相关多模态方法的性能比较" in p.text for p in check.paragraphs))


if __name__ == "__main__":
    main()
