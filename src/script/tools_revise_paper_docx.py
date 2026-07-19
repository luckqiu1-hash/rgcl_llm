from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


SRC = Path(r"D:\pycharm\rgcl_llm\paper_revision_work\paper_repaired_working.docx")
OUT = Path(r"C:\Users\hasee\Desktop\paper_revised_sec_rgcl.docx")


def set_text(p, text):
    p.clear()
    if text:
        p.add_run(text)


def set_style(p, style):
    try:
        p.style = style
    except Exception:
        pass


def insert_before(anchor, text, style=None):
    p = anchor.insert_paragraph_before(text)
    if style:
        set_style(p, style)
    return p


def main():
    doc = Document(str(SRC))
    p = doc.paragraphs

    set_text(p[0], "面向中文有害 Meme 检测的解释增强反事实检索对比学习方法")
    set_text(p[1], "Semantic Explanation and Counterfactual Evidence Enhanced Retrieval-Guided Contrastive Learning for Chinese Harmful Meme Detection")
    set_text(p[3], "作者信息待补充")
    set_text(p[4], "")
    set_text(p[5], "")

    set_text(p[10], "随着社交媒体和生成式大模型的快速发展，图像与文本共同构成的 Meme 已成为网络内容传播的重要载体，也使有害内容呈现出更强的隐蔽性和跨模态性。中文有害 Meme 往往依赖谐音、网络用语、文化典故以及图文组合后的隐含语义来表达侮辱、歧视或攻击意图，给自动内容审核带来挑战。现有 CLIP 类视觉语言模型和检索引导对比学习方法能够在一定程度上缓解混淆样本问题，但仍主要依赖原始图像与 OCR 文本的隐式表示，难以稳定捕获中文语境中的深层语义证据。")
    set_text(p[11], "针对上述问题，本文提出 SEC-RGCL，一种语义解释与反事实证据增强的检索引导对比学习框架。该方法首先利用多模态大模型生成中性的语义解释句，将图像、OCR 文本及其组合语义显式化；随后通过上下文感知门控、FiLM 调制和语义残差连接，将解释信息以受控方式注入图文联合表示，避免解释噪声直接主导分类结果。在此基础上，本文进一步设计反事实证据学习模块，根据分类器权重和样本表示激活估计关键维度，并构造证据削弱后的反事实表示，以约束模型依赖稳定且具有判别意义的语义证据。")
    set_text(p[12], "本文在 ToxiCN MM 中文有害 Meme 数据集上进行实验，并以 RGCL 为主要基线开展整体性能比较、模块消融、辅助文本对比和反事实响应分析。实验结果表明，SEC-RGCL 在 Acc 和 AUC 上分别达到 0.8392 和 0.8622，较原始 RGCL 基线提升 2.75 和 2.71 个百分点。进一步分析显示，语义解释能够提供比重复输入 OCR 文本更有效的辅助语义信息，反事实证据学习则增强了模型对关键判别维度的响应一致性。")

    set_text(p[13], "1 引言")
    set_style(p[13], "Head1")
    set_text(p[14], "随着社交媒体平台的快速发展，网络空间中的有害内容逐渐呈现出规模化、隐蔽化和多模态化的传播趋势。Meme 通常由图像和文字共同构成，能够以幽默、讽刺或隐喻的方式快速传播观点。与纯文本有害言论不同，有害 Meme 的攻击性含义往往并不完全体现在单一模态中，而是由视觉元素、OCR 文本以及二者之间的语用关系共同触发。因此，构建能够理解图文组合语义的有害 Meme 自动检测方法，对于在线社区治理和内容安全具有重要意义。")
    set_text(p[16], "图 1. 中文有害 Meme 中隐式表达与显式表达的示例。")
    set_text(p[17], "准确检测有害 Meme 的核心困难之一是混淆样本（confounder memes）。这类样本在图像或文字层面可能高度相似，但由于图文组合关系不同，其整体语义和类别标签会发生显著变化。例如，一个有害 Meme 可能与良性 Meme 共享相同文本，仅因配图不同而表达完全不同的含义；也可能共享相同图像，仅因文字变化而从攻击性表达转变为无害表达。由于这类样本在视觉语言表示空间中容易被映射到相近位置，即使是较强的 CLIP 类检测系统也可能对细粒度语义差异不够敏感。")
    set_text(p[18], "现有研究主要从两条路径改进该问题：一类方法设计跨模态融合模块，以增强图像与文本之间的对齐；另一类方法通过检索增强对比学习，动态选择同类伪正样本和异类困难负样本，从而学习更具判别性的表示空间。然而，对于中文有害 Meme 而言，许多关键线索来自谐音、典故、网络亚文化和语境暗示。仅依赖原始图像与 OCR 文本的隐式编码，模型仍可能受到浅层相似性、数据偏差或偶然相关特征的影响，难以稳定捕获真正支持有害性判断的语义证据。")
    set_text(p[19], "为此，本文提出 SEC-RGCL，将显式语义解释和反事实证据约束引入检索引导对比学习框架。与直接让大模型输出有害/无害标签不同，本文要求多模态大模型生成一条中性英文解释句，用于描述图像与 OCR 文本共同传达的具体含义，而不包含道德判断或分类标签。该解释句作为额外语义信号参与表示学习，使模型能够获得原始编码器难以显式表达的文化背景、隐喻关系和图文组合语义。")
    set_text(p[20], "进一步地，本文引入归因引导的反事实证据学习机制，以检验并约束模型是否依赖有效的判别证据。具体而言，模型根据最终表示维度的激活值和分类头权重估计关键证据维度，并对这些维度进行削弱以构造反事实表示。对于有害样本，若关键维度确实承载了有害性证据，则削弱后有害类别置信度应下降；对于无害样本，同类扰动不应异常提高有害类别置信度。该约束有助于减少模型对局部高频词、视觉 shortcut 或伪相关模式的依赖。")
    set_text(p[21], "本文以 RGCL 为基础骨架，在语义增强后的表示空间中执行检索与对比学习。语义解释并不替代原始图文特征，而是通过门控、调制和残差机制对图文表示进行补充与校准；检索机制则继续承担拉近同类样本、拉远异类混淆样本的作用。通过二者结合，模型既保留 RGCL 对混淆样本的判别优势，又获得面向中文语境的显式语义补充。")
    set_text(p[22], "本文在 ToxiCN MM 数据集上验证所提出方法的有效性，并围绕整体性能、模块消融、辅助文本类型和反事实响应进行分析。实验结果显示，语义解释增强和反事实证据学习均能提升 RGCL 基线性能，且二者联合使用时取得最佳结果。")
    set_text(p[23], "本文的主要贡献如下：")
    set_text(p[24], "（1）提出一种面向中文有害 Meme 检测的语义解释增强框架。该框架在原始图像与 OCR 文本之外引入中性语义解释，并通过上下文感知门控、FiLM 调制和语义残差连接进行受控融合，从而增强模型对隐式图文组合语义的建模能力。")
    set_text(p[25], "（2）提出一种归因引导的反事实证据学习机制。该机制在最终表示空间中削弱高贡献维度，并约束原始预测与反事实预测之间的关系，以提升模型对关键判别证据的依赖稳定性。")
    set_text(p[26], "（3）在 ToxiCN MM 数据集上开展系统实验。整体性能、消融实验、OCR 文本对比和反事实响应分析共同表明，SEC-RGCL 能够在原始 RGCL 基线之上取得稳定提升，并改善模型在隐式语义场景下的判别能力。")
    set_text(p[27], "")

    # Insert a missing Related Work section before the method section.
    anchor = p[28]
    related = [
        ("2 相关工作", "Head2"),
        ("2.1 多模态有害 Meme 检测", "Heading 2"),
        ("有害 Meme 检测要求模型同时理解视觉内容、图中文字以及二者组合后的语义。早期方法通常基于目标检测器或视觉语言预训练模型提取多模态特征，并通过跨模态注意力、特征拼接或逐维交互完成分类。近年来，CLIP 类模型因端到端结构简洁、图文对齐能力较强而被广泛用于 Meme 检测任务。然而，已有研究表明，CLIP 表示空间对 confounder memes 中细微但关键的图文差异仍不够敏感，容易将表层相似但语义标签相反的样本映射到邻近区域。", "Normal (Web)"),
        ("2.2 解释增强与中文文化语境", "Heading 2"),
        ("中文有害 Meme 的语义理解高度依赖文化背景、谐音、网络流行语和隐喻表达。参考论文中的 FG-E2HMD 强调，先生成文化感知解释再进行分类，有助于提升检测性能和决策透明度。本文借鉴这一“先解释、再利用解释增强判断”的思想，但不直接使用解释生成标签，而是将中性解释句作为表征学习的辅助语义信号，从而降低标签泄漏和过度推断风险。", "Normal (Web)"),
        ("2.3 检索引导对比学习与反事实约束", "Heading 2"),
        ("RGCL 通过动态检索伪正样本和困难负样本，构造 hatefulness-aware embedding space，是本文的主要 baseline。该方法能够缓解混淆样本在表示空间中过度接近的问题，但其检索仍主要依赖原始图文表示。本文在 RGCL 基础上引入语义解释增强的检索空间，并进一步加入反事实证据学习，使模型不仅学习更好的类间结构，也受到关键证据依赖关系的显式约束。", "Normal (Web)"),
    ]
    for text, style in related:
        insert_before(anchor, text, style)

    set_text(p[28], "3 方法")
    set_style(p[28], "Head2")
    set_text(p[29], "本文提出语义解释与反事实证据增强的检索引导对比学习框架（Semantic Explanation and Counterfactual Evidence Enhanced RGCL，SEC-RGCL）。该框架以 RGCL 的图文联合表示和动态检索机制为基础，加入语义解释生成与受控融合模块，并通过反事实证据学习约束模型对关键判别维度的依赖。")
    set_text(p[30], "原始 RGCL 通过动态检索 pseudo-gold positive 和 hard negative，在图文联合表征空间中拉近同类样本、拉远异类混淆样本，从而学习具有有害性判别能力的表示空间。该机制适合处理图像或文本相似但标签相反的混淆样本。然而，在中文有害 Meme 场景下，关键语义常来自网络用语、文化典故或图文组合后的隐含含义，仅依赖原始图文 embedding 可能无法提供足够稳定的判别依据。")
    set_text(p[31], "SEC-RGCL 的核心思想是：首先将 Meme 的图文组合语义显式化为一条中性解释句，再将该解释作为辅助语义信号注入图文表示；随后在语义增强表示空间中执行检索引导对比学习，并通过反事实扰动约束模型的证据依赖关系。语义解释模块主要回答“模型可以获得什么额外语义信息”，反事实证据模块则进一步回答“模型是否真正依赖这些有效证据进行判断”。")
    set_text(p[32], "下面分别介绍语义解释生成与编码、解释感知投影与融合、检索引导对比学习以及反事实证据学习。")
    set_text(p[35], "3.1 语义解释生成与编码")
    set_style(p[35], "Heading 2")
    set_text(p[41], "本文中的语义解释并非用于直接输出 harmful 或 non-harmful 的最终判断，而是被设计为一条用于 meme embedding 的中性英文语义解释句。该解释句概括 Visual + OCR 共同传达的具体含义，重点包括图像中的关键实体、OCR 文本的字面含义以及二者组合后的语义关系。在证据充分时，解释句还可以包含中文网络用语、谐音、文化引用、符号和群体指代等信息。与标签式解释不同，该设计能够降低标签泄漏风险，使解释信息更适合作为后续融合、检索和对比学习的语义补充。")
    set_text(p[44], "Generate one neutral English semantic explanation sentence to supplement image-text representations for meme embedding.\n\nGiven the meme image and OCR text, state the concrete meaning conveyed by the combination of Visual + OCR. Include relevant visual entities, OCR meaning, Chinese internet slang, puns, cultural references, symbols, events, groups, and outcomes only when they are clearly supported. If the implied meaning is unclear, describe only the literal visual content and literal OCR meaning. Do not output a harmful/harmless label or any moral judgment. Avoid words such as hateful, toxic, harmful, harmless, racist, offensive, discriminatory, or abusive. If a person, group, or social category is explicitly involved, describe it neutrally as being referenced, portrayed, compared, or associated, rather than attacked or targeted. Use one English sentence only.\n\nOCR: {}")
    set_text(p[49], "3.2 解释感知投影与融合")
    set_text(p[73], "3.3 语义增强的检索引导对比学习")
    set_text(p[87], "其中，检索得到的同标签近邻作为 pseudo-gold positive，异标签近邻作为 hard negative，τ 表示温度系数。与仅基于原始图文 embedding 的检索方式相比，语义增强后的检索空间能够更充分地刻画图文组合语义，使伪正样本更可能共享相近的判别依据，困难负样本也更可能对应真正具有迷惑性的混淆样本。因此，对比学习执行的拉近与拉远过程不再只依赖图像或文字的表层相似性，而是建立在更明确的语义补充之上。")
    set_text(p[88], "3.4 反事实证据学习")

    set_text(p[112], "4 实验与分析")
    set_style(p[112], "Head2")
    set_text(p[113], "本章对 SEC-RGCL 进行实验评估。首先介绍数据集、模型设置、评价指标和参数设置；随后比较 SEC-RGCL 与原始 RGCL 基线的整体性能，并通过模块消融分析语义解释增强与反事实证据学习的独立贡献；最后进一步考察语义解释相对于原始 OCR 文本的作用，以及反事实证据学习对模型内部响应的影响。")
    set_text(p[114], "4.1 实验设置")
    set_style(p[114], "Heading 2")
    set_text(p[115], "4.1.1 数据集与任务")
    set_style(p[115], "Heading 2")
    set_text(p[118], "4.1.2 模型设置")
    set_style(p[118], "Heading 2")
    set_text(p[121], "4.2 评价指标")
    set_style(p[121], "Heading 2")
    set_text(p[125], "4.3 参数设置")
    set_style(p[125], "Heading 2")
    set_text(p[130], "作者待补充：为保证实验可复现，最终版本仍需补充图像编码器、OCR 文本编码器、解释编码器、多模态大模型及版本、优化器、学习率、batch size、温度系数、损失权重、top-k、削弱比例、margin、随机种子、检索库更新频率和硬件环境等细节。")
    set_text(p[131], "4.4 与基线模型的性能比较")
    set_style(p[131], "Heading 2")
    set_text(p[134], "与 Baseline 相比，SEC-RGCL 的 Acc 由 0.8117 提升至 0.8392，提高 2.75 个百分点；AUC 由 0.8351 提升至 0.8622，提高 2.71 个百分点。两项指标上的一致增益表明，在原始 RGCL 框架中同时引入语义解释增强与反事实证据学习，能够有效改善中文有害 Meme 的分类性能。")
    set_text(p[135], "4.5 模块消融实验")
    set_style(p[135], "Heading 2")
    set_text(p[140], "完整 SEC-RGCL 在两项指标上均取得最优结果。相较于 Baseline + Exp，SEC-RGCL 的 Acc 和 AUC 进一步提高 1.17 和 0.96 个百分点；相较于 Baseline + CF，分别提高 2.17 和 2.06 个百分点。上述结果表明，语义解释增强与反事实证据学习具有互补性：前者扩展模型可利用的显式语义信息，后者约束模型对判别证据的依赖方式，二者结合能够带来更稳定的性能收益。")
    set_text(p[141], "4.6 语义解释信息作用分析")
    set_style(p[141], "Heading 2")
    set_text(p[144], "Baseline + Text 的 Acc 和 AUC 分别为 0.8217 和 0.8445，较 Baseline 提高 1.00 和 0.94 个百分点，表明增加辅助文本分支能够带来一定性能增益。使用大模型语义解释后，Baseline + Exp 的 Acc 和 AUC 进一步达到 0.8275 和 0.8526，分别高于 Baseline + Text 0.58 和 0.81 个百分点。")
    set_text(p[145], "由于 Baseline + Text 与 Baseline + Exp 均保留额外文本编码和融合路径，二者之间的性能差异主要反映辅助文本内容本身的作用。相较于重复输入原始 OCR 文本，语义解释能够概括图像与文字共同表达的含义，并补充原始文本中未显式表达的语义关系，因此为分类模型提供了更有效的辅助信息。需要强调的是，该实验评价的是解释作为模型输入的贡献，而不是解释文本本身的人类可读质量。")
    set_text(p[146], "4.7 反事实证据机制分析")
    set_style(p[146], "Heading 2")
    set_text(p[151], "实验在同一组 370 个有害样本上比较完整模型 full 与移除 CF 模块的 no_cf。Mean Drop 和 Median Drop 分别表示置信度下降量的均值与中位数；Drop Rate 表示置信度出现下降的样本比例；Toxic→Non-toxic Flip 表示干预后预测由有害类别翻转为无害类别的样本比例。结果如表 4-5 所示。")
    set_text(p[154], "full 的 Drop Rate 为 0.941，即 94.1% 的有害样本在干预后出现置信度下降，而 no_cf 的该比例为 61.6%。此外，full 的 Toxic→Non-toxic Flip 为 31.4%，no_cf 仅为 0.3%。由于上述指标均在定向削弱高贡献维度后计算，较高的置信度下降和类别翻转并不表示模型在正常输入下性能下降，而是表明完整模型对关键证据维度具有更强、更一致的响应。")
    set_text(p[156], "4.8 本章小结")
    set_style(p[156], "Heading 2")

    # Add case study section before the first case image/text block.
    insert_before(p[159], "5 案例分析与讨论", "Head2")
    insert_before(p[159], "为进一步说明 SEC-RGCL 的作用，本文选取两个原始 RGCL 与通用大模型均容易误判的样本进行分析。案例分析的目的不是替代定量实验，而是展示语义解释如何补充图文组合含义，以及反事实证据约束如何帮助模型减少对表层线索的依赖。", "Normal (Web)")
    set_text(p[160], "案例 1 中，Meme 使用“垃圾袋”这一视觉隐喻讽刺对方“能装”，其攻击性并不完全来自 OCR 文本本身，而是来自图像符号与中文网络语义之间的组合关系。语义解释将该隐含关系显式化后，SEC-RGCL 能够更好地区分表面无害表达与实际讽刺意图。")
    set_text(p[161], "预测结果：RGCL 错误；Qwen-3.5 错误；SEC-RGCL 正确。")
    set_text(p[164], "案例 2 中，亲吻场景与“科学家正在努力治愈她们”的文字共同构成对性少数群体的病理化隐喻。该样本的有害性来自图文组合后的社会语义，而非单一词语触发。SEC-RGCL 通过解释增强获得更明确的语义补充，因此能够识别原始图文表示中不易捕获的歧视性含义。")
    set_text(p[165], "预测结果：RGCL 错误；Qwen-3.5 错误；SEC-RGCL 正确。")

    insert_before(p[166], "6 结论", "Head2")
    insert_before(p[166], "本文提出 SEC-RGCL，一种面向中文有害 Meme 检测的语义解释与反事实证据增强检索对比学习方法。该方法在 RGCL 的动态检索对比学习框架基础上，引入多模态大模型生成的中性语义解释，并通过门控、FiLM 调制和残差连接将解释信息受控注入图文表示。同时，本文设计反事实证据学习模块，通过削弱高贡献维度并约束预测响应，促使模型依赖更稳定的判别证据。", "Normal (Web)")
    insert_before(p[166], "在 ToxiCN MM 数据集上的实验表明，SEC-RGCL 在 Acc 和 AUC 上均优于原始 RGCL。消融实验进一步验证了语义解释增强和反事实证据学习的独立贡献；辅助文本对比说明，语义解释相较于重复输入 OCR 文本能够提供更有效的补充信息；反事实响应分析则表明，完整模型对关键证据维度具有更一致的预测变化。", "Normal (Web)")
    insert_before(p[166], "本文仍存在若干限制。首先，语义解释质量依赖底层多模态大模型，当模型缺乏特定文化知识或产生过度推断时，解释可能引入噪声。其次，反事实证据学习在表示维度上进行干预，不能直接等同于人类可解释的语义因果关系。未来工作可进一步研究可验证的解释生成、跨数据集泛化以及更细粒度的人类可读证据定位。", "Normal (Web)")

    set_text(p[167], "作者将在最终版本中补充致谢信息。")
    refs = [
        "[1]\tJingbiao Mei, Jinghong Chen, Weizhe Lin, Bill Byrne, and Marcus Tomalin. 2024. Improving Hateful Meme Detection through Retrieval-Guided Contrastive Learning. In Proceedings of the 62nd Annual Meeting of the Association for Computational Linguistics (Volume 1: Long Papers), 5333-5347.",
        "[2]\tXinhao Chen, Dongxin Wen, and Decheng Zuo. 2025. Towards Detecting Chinese Harmful Memes with Fine-Grained Explanatory Augmentation. Electronics 14(17):3504.",
        "[3]\tDouwe Kiela, Hamed Firooz, Aravind Mohan, Vedanuj Goswami, Amanpreet Singh, Pratik Ringshia, and Davide Testuggine. 2020. The Hateful Memes Challenge: Detecting Hate Speech in Multimodal Memes. In Advances in Neural Information Processing Systems.",
        "[4]\tNithish Kumar and Radhakrishnan Nandakumar. 2022. Hate-CLIPper: Multimodal Hateful Meme Classification Based on Cross-modal Interaction of CLIP Features.",
        "[5]\tAlec Radford, Jong Wook Kim, Chris Hallacy, Aditya Ramesh, Gabriel Goh, Sandhini Agarwal, Girish Sastry, Amanda Askell, Pamela Mishkin, Jack Clark, Gretchen Krueger, and Ilya Sutskever. 2021. Learning Transferable Visual Models From Natural Language Supervision. In Proceedings of ICML.",
    ]
    for idx, text in zip(range(169, 174), refs):
        set_text(p[idx], text)
        set_style(p[idx], "Bib_entry")
    for idx in range(174, 184):
        set_text(p[idx], "")

    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
