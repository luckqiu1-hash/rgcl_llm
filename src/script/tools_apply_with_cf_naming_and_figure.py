from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


SRC = Path(r"C:\Users\hasee\Desktop\paper_semantic_refocused_sec_rgcl.docx")
OUT = Path(r"C:\Users\hasee\Desktop\paper_semantic_refocused_with_cf_academic_figure.docx")
FIG = Path(r"D:\pycharm\rgcl_llm\src\cf_evidence_behavior\cf_toxic_only_paper_academic.png")


def set_text(p, text):
    p.clear()
    p.add_run(text)


def replace_text_in_paragraphs(doc):
    replacements = {
        "实验在同一组 370 个有害样本上比较完整模型 full 与移除 CF 模块的 no_cf。Mean Drop 和 Median Drop 分别表示置信度下降量的均值与中位数；Drop Rate 表示置信度出现下降的样本比例；Toxic→Non-toxic Flip 表示干预后预测由有害类别翻转为无害类别的样本比例。结果如表 4-5 所示。":
            "实验在同一组 370 个有害样本上比较加入反事实证据学习的模型 with_cf 与移除该模块的模型 no_cf。Mean Drop 和 Median Drop 分别表示置信度下降量的均值与中位数；Drop Rate 表示置信度出现下降的样本比例；Toxic→Non-toxic Flip 表示干预后预测由有害类别翻转为无害类别的样本比例。结果如表 4-5 和图 3 所示。",
        "在关键维度被削弱后，full 的 Mean Drop 为 0.149，明显高于 no_cf 的 0.009；其 Median Drop 也由 0.067 提高至 0.109。这说明加入 CF 模块后，模型所识别的高贡献维度与有害类别预测之间具有更强的对应关系。":
            "在关键维度被削弱后，with_cf 的 Mean Drop 为 0.149，明显高于 no_cf 的 0.009；其 Median Drop 也由 0.067 提高至 0.109。这说明加入 CF 模块后，模型所识别的高贡献维度与有害类别预测之间具有更强的对应关系。",
        "full 的 Drop Rate 为 0.941，即 94.1% 的有害样本在干预后出现置信度下降，而 no_cf 的该比例为 61.6%。此外，full 的 Toxic→Non-toxic Flip 为 31.4%，no_cf 仅为 0.3%。由于上述指标均在定向削弱高贡献维度后计算，较高的置信度下降和类别翻转并不表示模型在正常输入下性能下降，而是表明完整模型对关键证据维度具有更强、更一致的响应。":
            "with_cf 的 Drop Rate 为 0.941，即 94.1% 的有害样本在干预后出现置信度下降，而 no_cf 的该比例为 61.6%。此外，with_cf 的 Toxic→Non-toxic Flip 为 31.4%，no_cf 仅为 0.3%。由于上述指标均在定向削弱高贡献维度后计算，较高的置信度下降和类别翻转并不表示模型在正常输入下性能下降，而是表明加入 CF 后模型对关键证据维度具有更强、更一致的响应。",
        "结合消融实验结果，CF 模块不仅提高了 Acc 和 AUC，也增强了模型在关键证据受到干预时的响应一致性。需要说明的是，该分析刻画的是模型内部高贡献维度与预测结果之间的关联，不能据此直接断言这些维度具有明确的人类语义解释。换言之，反事实证据学习提供的是一种表示层面的稳定性约束，而不是完整的因果解释。":
            "结合消融实验结果，CF 模块不仅提高了 Acc 和 AUC，也增强了模型在关键证据受到干预时的响应一致性。需要说明的是，该分析刻画的是模型内部高贡献维度与预测结果之间的关联，不能据此直接断言这些维度具有明确的人类语义解释。换言之，反事实证据学习提供的是一种表示层面的稳定性约束，而不是完整的因果解释。",
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in replacements:
            set_text(p, replacements[text])
        elif "full" in text:
            set_text(p, text.replace("full", "with_cf"))


def replace_text_in_tables(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text == "full":
                        set_text(p, "with_cf")
                    elif "full" in text:
                        set_text(p, text.replace("full", "with_cf"))


def insert_paragraph_before(anchor, text="", style=None):
    p = anchor.insert_paragraph_before(text)
    if style:
        try:
            p.style = style
        except Exception:
            pass
    return p


def insert_figure(doc):
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("在关键维度被削弱后，with_cf 的 Mean Drop"):
            anchor = p
            break
    if anchor is None:
        raise ValueError("Could not find insertion anchor after Table 4-5.")

    fig_p = insert_paragraph_before(anchor)
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_p.add_run().add_picture(str(FIG), width=Inches(6.3))

    cap = insert_paragraph_before(
        anchor,
        "图 3. 反事实证据削弱下的有害类别置信度响应。左图展示 with_cf 与 no_cf 在有害样本上的 Drop 分布，其中 Drop = p(toxic | original) − p(toxic | counterfactual)；右图汇总 Mean Drop、Drop Rate 和 Toxic→Non-toxic Flip。更高的下降幅度和翻转比例表明模型对被选中关键证据维度具有更强响应。",
        "Caption",
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document(str(SRC))
    replace_text_in_paragraphs(doc)
    replace_text_in_tables(doc)
    insert_figure(doc)
    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
