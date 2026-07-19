from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


OUT = Path(r"C:\Users\hasee\Desktop\SEC-RGCL_3.2_fusion_revision.docx")


PARAGRAPHS = [
    (
        "说明",
        "以下内容用于替换论文第 3.2 节“解释感知投影与融合”中从“在获得图像、OCR 文本和语义解释三路信息后”开始的融合机制描述。该版本按照代码中的真实实现改写：基础图文表示采用拼接而非逐维相乘；语义解释先经过适配与门控；随后通过 MODA 语义引导交互、FiLM 有界校准、语义残差补充和层级加权组合得到最终 embedding。",
    ),
    (
        "正文",
        "在获得图像、OCR 文本和语义解释三路信息后，本文首先将不同来源的特征映射到统一隐空间。给定第 i 个样本的图像、OCR 文本和语义解释，分别记为 I_i、T_i 和 E_i，模型通过对应编码器获得初始特征：",
    ),
    ("formula", "v_i = F_v(I_i),    t_i = F_t(T_i),    e_i = F_e(E_i),"),
    (
        "正文",
        "其中，v_i、t_i 和 e_i 分别表示图像特征、OCR 文本特征和语义解释特征。由于三类特征来自不同编码空间，本文进一步使用线性投影和归一化操作将其映射到相同维度：",
    ),
    ("formula", "v~_i = LN(W_v v_i + b_v),"),
    ("formula", "t~_i = LN(W_t t_i + b_t),"),
    ("formula", "e~_i = LN(W_e e_i + b_e)."),
    (
        "正文",
        "图像和 OCR 文本构成模型的基础多模态证据。本文首先将二者拼接并归一化，得到基础图文表示：",
    ),
    ("formula", "x_i^0 = LN([v~_i ; t~_i])."),
    (
        "正文",
        "该表示保留了原始视觉内容和 OCR 文本信息，是后续分类与检索表示学习的主干。",
    ),
    (
        "正文",
        "语义解释能够补充原始图文编码中难以显式表达的隐含含义，但生成式解释也可能包含冗余或与当前样本弱相关的信息。因此，本文不直接将解释特征作为独立判别输入，而是先通过适配层和门控机制对解释信息进行筛选：",
    ),
    ("formula", "h_i^e = A(e~_i),"),
    ("formula", "alpha_i = sigmoid(W_g [v~_i ; t~_i ; h_i^e] + b_g),"),
    ("formula", "e^_i = alpha_i ⊙ h_i^e,"),
    (
        "正文",
        "其中，A(·) 表示解释适配层，sigmoid(·) 表示 Sigmoid 函数，alpha_i 为解释特征的维度级门控权重。该过程使模型能够根据当前图像和文本内容动态选择更相关的解释信息。",
    ),
    (
        "正文",
        "在获得门控后的解释表示 e^_i 后，模型采用渐进式方式将其注入图文表示。首先，本文使用语义引导的跨模态交互模块建模图像、OCR 文本和解释信息之间的关系：",
    ),
    ("formula", "z_i = M(v~_i, t~_i, e^_i),"),
    ("formula", "x_i^base = LN(x_i^0 + tanh(lambda_m) z_i),"),
    (
        "正文",
        "其中，M(·) 表示语义引导的跨模态融合模块，lambda_m 为可学习的残差系数。该步骤使解释信息参与图文交互建模，同时通过残差形式保留原始图文表示结构。",
    ),
    (
        "正文",
        "随后，本文利用 FiLM 调制对基础表示进行特征维度校准。解释表示被用于生成缩放因子和偏置项：",
    ),
    ("formula", "[gamma_i ; beta_i] = W_f e^_i + b_f,"),
    ("formula", "gamma_i = 0.05 tanh(gamma_i),    beta_i = 0.05 tanh(beta_i)."),
    (
        "正文",
        "基于上述参数，图文表示被更新为：",
    ),
    ("formula", "x_i^film = x_i^base + tanh(lambda_f)(gamma_i ⊙ x_i^base + beta_i),"),
    (
        "正文",
        "其中，lambda_f 为可学习调制强度。通过对 gamma_i 和 beta_i 进行幅度约束，模型能够利用解释信息校准图文表示，同时避免生成解释对基础表示产生过强扰动。",
    ),
    (
        "正文",
        "进一步地，本文引入语义残差连接，为图文表示补充解释中包含的高层语义线索：",
    ),
    ("formula", "r_i^e = 0.05 tanh(W_r e^_i + b_r),"),
    ("formula", "eta_i = sigmoid(W_s [v~_i ; t~_i ; e^_i] + b_s),"),
    ("formula", "x_i^sem = x_i^film + tanh(lambda_s)(eta_i ⊙ r_i^e),"),
    (
        "正文",
        "其中，eta_i 为语义残差门控权重，lambda_s 为可学习残差强度。该设计使解释信息以轻量残差形式补充隐含语义，而不是替代原始图文证据。",
    ),
    (
        "正文",
        "最后，模型对不同阶段的表示进行归一化加权组合：",
    ),
    ("formula", "[omega_0, omega_1, omega_2] = softmax(q),"),
    ("formula", "x_i = LN(omega_0 LN(x_i^base) + omega_1 LN(x_i^film) + omega_2 LN(x_i^sem))."),
    (
        "正文",
        "其中，q 为可学习的层级权重。最终表示 x_i 输入残差分类头 f_m，得到用于分类和检索对比学习的 embedding：",
    ),
    ("formula", "g_i = f_m(x_i)."),
    (
        "正文",
        "该融合过程以原始图像和 OCR 文本表示为主干，将语义解释作为条件调制和残差补充信号逐步引入模型。一方面，门控和幅度约束降低了生成式解释中冗余信息对表示空间的干扰；另一方面，FiLM 调制和语义残差为图文表示补充了关于图文组合含义、隐喻关系和语境指代的显式线索，从而增强模型对中文有害 Meme 隐含语义的建模能力。",
    ),
]


def set_cell_text_font(paragraph, east_asia="宋体", latin="Times New Roman", size=12):
    for run in paragraph.runs:
        run.font.name = latin
        run.font.size = Pt(size)
        rpr = run._element.get_or_add_rPr()
        fonts = rpr.rFonts
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.append(fonts)
        fonts.set(qn("w:eastAsia"), east_asia)
        fonts.set(qn("w:ascii"), latin)
        fonts.set(qn("w:hAnsi"), latin)


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)
    return p


def main():
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Heading 1"].font.name = "黑体"
    styles["Heading 2"].font.name = "黑体"

    title = doc.add_heading("SEC-RGCL 第 3.2 节融合机制改写稿", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    current_heading = None
    for kind, text in PARAGRAPHS:
        if kind in {"说明", "正文"} and kind != current_heading:
            if kind == "说明":
                doc.add_heading("使用说明", level=2)
            elif kind == "正文":
                doc.add_heading("可替换正文", level=2)
            current_heading = kind

        if kind == "formula":
            add_formula(doc, text)
        else:
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.line_spacing = 1.5
            set_cell_text_font(p)

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
