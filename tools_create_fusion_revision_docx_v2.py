from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


OUT = Path(r"C:\Users\hasee\Desktop\SEC-RGCL_3.2_fusion_revision_v6_integrated.docx")


CONTENT = [
    (
        "说明",
        "本版本用于替换论文第 3.2 节“解释感知投影与融合”中关于语义解释融合机制的描述。相较上一版，本版本进一步明确了 A(·) 和 M(·) 的具体含义：A(·) 对应代码中的解释适配器，由线性层、SiLU 激活和 Dropout 组成；M(·) 不再使用特定模块名称描述，而是改写为由线性 token 化、图文关系 token、多头注意力交互、语义引导注意力、池化和线性投影构成的跨模态融合过程，避免读者因不熟悉特定模块名而产生理解障碍。",
    ),
    (
        "正文",
        "在获得图像、OCR 文本和语义解释三路信息后，本文首先将不同来源的特征映射到统一隐空间。给定第 i 个样本的图像、OCR 文本和语义解释，分别记为 I_i、T_i 和 E_i，模型通过对应编码器获得初始特征：",
    ),
    ("formula", "v_i = F_v(I_i),    t_i = F_t(T_i),    e_i = F_e(E_i),"),
    (
        "正文",
        "其中，v_i、t_i 和 e_i 分别表示图像特征、OCR 文本特征和语义解释特征。由于三类特征来自不同编码空间，本文进一步使用线性投影和归一化操作将其映射到相同维度：",
    ),
    ("formula", "v~_i = LN(W_v v_i + b_v),"),
    ("formula", "t~_i = LN(W_t t_i + b_t),"),
    ("formula", "e~_i = LN(W_e e_i + b_e)."),
    (
        "正文",
        "图像和 OCR 文本是有害 Meme 检测的基础证据来源。本文首先将二者拼接并归一化，得到基础图文表示：",
    ),
    ("formula", "x_i^0 = LN([v~_i ; t~_i])."),
    (
        "正文",
        "该表示保留了原始视觉内容和 OCR 文本信息，并作为后续分类与检索对比学习的主干表示。相比直接将语义解释与图文特征拼接，以图文表示作为主干能够更好地保持 RGCL 原有表示空间中的图文相似性结构。",
    ),
    (
        "正文",
        "语义解释能够补充图像和 OCR 文本中难以显式表达的隐含含义，例如图文组合后的语用关系、隐喻指向和语境暗示。然而，语义解释由生成式模型产生，其中可能包含冗余描述或与当前样本弱相关的信息。如果这些信息被无差别地引入表示空间，模型可能受到解释噪声影响。为此，本文首先在解释特征进入融合过程前引入解释筛选门控：",
    ),
    ("formula", "h_i^e = A(e~_i) = Dropout(SiLU(W_a e~_i + b_a)),"),
    ("formula", "alpha_i = sigmoid(W_g [v~_i ; t~_i ; h_i^e] + b_g),"),
    ("formula", "e^_i = alpha_i ⊙ h_i^e,"),
    (
        "正文",
        "其中，A(·) 表示解释适配器，对应线性变换、SiLU 激活和 Dropout 的组合，用于将解释特征转换为更适合后续融合的上下文表示；sigmoid(·) 表示 Sigmoid 函数，alpha_i 为维度级解释筛选权重。该门控的作用不是完成最终融合，而是在解释信息进入后续交互前，根据当前图像和 OCR 文本内容筛选更相关的解释维度，从源头降低冗余解释对表示学习的干扰。",
    ),
    (
        "正文",
        "在获得筛选后的解释表示 e^_i 后，模型首先利用语义引导的跨模态交互模块建模图像、OCR 文本和解释信息之间的关系：",
    ),
    ("formula", "z_i = M(v~_i, t~_i, e^_i),"),
    ("formula", "x_i^base = LN(x_i^0 + tanh(lambda_m) z_i),"),
    (
        "正文",
        "其中，M(·) 表示语义引导的跨模态融合过程。具体而言，该过程首先通过线性层将图像特征、OCR 文本特征和筛选后的解释特征分别展开为若干 token，使原本的全局向量能够以 token 序列形式参与注意力计算。为显式建模图文之间的关系，模型进一步构造两类关系 token：图文差异 token |v~_i - t~_i| 用于刻画图像与文本之间的不一致性，图文乘积 token v~_i ⊙ t~_i 用于刻画二者在特征维度上的对齐关系。随后，图像 token、文本 token、解释 token 和关系 token 被拼接为统一序列，并通过多头注意力进行跨模态交互，使视觉信息能够关注文本、解释和关系线索，文本与解释信息也能够根据视觉内容进行更新。在注意力交互之后，模型进一步以文本 token 和解释 token 作为语义引导信号，对整体 token 序列进行一次语义校准，使融合表示更关注与图文组合含义相关的特征。最后，模型分别对图像、文本、解释和关系 token 进行平均池化，并将池化结果拼接后输入线性投影层，得到语义引导的跨模态交互表示 z_i。lambda_m 为可学习的残差系数。该步骤使解释信息参与图文交互建模，同时通过残差形式保留基础图文表示，避免显式解释直接替代原始图文证据。",
    ),
    (
        "正文",
        "随后，本文使用 FiLM 调制对基础图文表示进行特征维度校准。FiLM 的作用是根据解释信息调整不同表示维度的响应强度，使模型能够突出与语义解释一致的图文特征，并抑制与当前含义弱相关的维度。具体而言，解释表示被用于生成缩放因子和偏置项：",
    ),
    ("formula", "[gamma_i ; beta_i] = W_f e^_i + b_f,"),
    ("formula", "gamma_i = 0.05 tanh(gamma_i),    beta_i = 0.05 tanh(beta_i)."),
    (
        "正文",
        "基于上述参数，图文表示被更新为：",
    ),
    ("formula", "x_i^film = x_i^base + tanh(lambda_f)(gamma_i ⊙ x_i^base + beta_i),"),
    (
        "正文",
        "其中，lambda_f 为可学习调制强度。由于 gamma_i 和 beta_i 均经过幅度约束，FiLM 分支对图文表示的影响被限制在较小范围内，从而实现条件校准而非大幅重写基础表示。",
    ),
    (
        "正文",
        "FiLM 调制主要作用于已有图文表示的维度响应，而部分隐含语义可能并未被原始编码器充分表达。为补充这类信息，本文进一步引入语义残差连接，将解释特征映射为轻量的残差补充项：",
    ),
    ("formula", "r_i^e = 0.05 tanh(W_r e^_i + b_r)."),
    (
        "正文",
        "在残差注入阶段，本文使用残差注入权重控制补充项对已校准表示的影响强度：",
    ),
    ("formula", "eta_i = sigmoid(W_s [v~_i ; t~_i ; e^_i] + b_s),"),
    ("formula", "x_i^sem = x_i^film + tanh(lambda_s)(eta_i ⊙ r_i^e),"),
    (
        "正文",
        "其中，eta_i 为语义残差注入权重，lambda_s 为可学习残差强度。与前述解释筛选门控不同，eta_i 不再用于选择解释内容本身，而是用于控制语义残差对当前表示的注入幅度。这样，模型先筛选解释内容，再约束其残差影响强度，从而在利用显式语义线索的同时保持图文主干表示的稳定性。",
    ),
    (
        "正文",
        "最后，模型将基础表示、FiLM 调制表示和语义残差增强表示进行归一化层级加权组合：",
    ),
    ("formula", "[omega_0, omega_1, omega_2] = softmax(q),"),
    ("formula", "x_i = LN(omega_0 LN(x_i^base) + omega_1 LN(x_i^film) + omega_2 LN(x_i^sem))."),
    (
        "正文",
        "其中，q 为可学习的层级权重。该组合方式使模型能够在不同样本上自适应地平衡基础图文证据、解释调制后的表示和语义残差增强表示。最终表示 x_i 输入残差分类头 f_m，得到用于分类和检索对比学习的 embedding：",
    ),
    ("formula", "g_i = f_m(x_i)."),
    (
        "正文",
        "得到语义增强 embedding 后，本文沿用 RGCL 的检索引导对比学习框架，在该表示空间中构建 pseudo-gold positive 和 hard negative。具体而言，训练集中所有样本的 embedding 及其标签被存入检索数据库：",
    ),
    ("formula", "G = {(g_i, y_i)}_{i=1}^N."),
    (
        "正文",
        "对于第 i 个训练样本，模型在数据库中检索与其标签相同且相似度最高的样本作为 pseudo-gold positive，记其 embedding 为 u_i；同时检索标签不同但相似度最高的样本作为 hard negative，记为 r_i。当前 batch 中与样本 i 标签不同的样本构成 in-batch negative 集合：",
    ),
    ("formula", "B_i^- = {g_k | g_k in B, y_k != y_i}."),
    (
        "正文",
        "因此，第 i 个样本的负样本集合为：",
    ),
    ("formula", "H_i = {r_i} ∪ B_i^-."),
    (
        "正文",
        "本文使用余弦相似度 c(g_i, g_j) 度量样本间距离，并将检索引导对比损失写为：",
    ),
    ("formula", "L_{RGCL,i} = -log exp(c(g_i,u_i)/tau) / [exp(c(g_i,u_i)/tau) + Σ_{h in H_i} exp(c(g_i,h)/tau)]."),
    (
        "正文",
        "原始 RGCL 的核心优势在于通过检索同标签近邻和异标签困难样本来优化表示空间，但其检索依据主要来自图像和文本编码后的联合表示。在中文有害 Meme 场景下，许多样本的关键差异并不直接体现在图像对象或 OCR 字面文本中，而是体现在图文组合后的语用含义、谐音隐喻、网络用语和语境指代中。因此，仅基于原始图文 embedding 的近邻检索可能受到表层相似性的影响：部分同标签近邻虽然视觉或文本相似，却未必共享相同的判别语义；部分异标签困难样本虽然在表示空间中接近，也可能只是由模板、人物表情或常见文字形式造成的浅层相似。",
    ),
    (
        "正文",
        "为缓解这一问题，本文使用语义增强 embedding 构建检索空间。由于该表示已经融合了图文组合含义、隐含语义和语境线索，检索得到的同标签近邻更可能共享相近的判别依据，而异标签近邻也更可能对应真正具有迷惑性的 confounder 样本。换言之，语义解释并不是替代 RGCL 的检索对比机制，而是为其提供更符合中文 Meme 语义结构的表示空间，使对比学习中的拉近与拉远过程不再只依赖图像或文本的表层相似性，而是更多地围绕显式化后的语义证据展开。因此，语义解释的作用不仅体现在分类表示中，也体现在对比学习阶段的近邻选择和表示空间优化过程中。",
    ),
    (
        "正文",
        "综上，该融合过程以原始图像和 OCR 文本表示为主干，将语义解释作为辅助调制和残差补充信号逐步引入模型。解释筛选门控用于降低生成式解释中的冗余信息，FiLM 调制用于对图文表示进行有界校准，语义残差用于补充原始编码器难以显式捕获的高层语义线索。通过这种方式，模型能够在保持基础图文表示稳定性的同时增强对中文有害 Meme 隐含语义的建模能力。",
    ),
]


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=12):
    run.font.name = latin
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    set_run_font(run, east_asia="Cambria Math", latin="Cambria Math", size=11)
    return p


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    title = doc.add_heading("SEC-RGCL 第 3.2 节融合机制改写稿 V2", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    current = None
    for kind, text in CONTENT:
        if kind == "说明" and current != "说明":
            doc.add_heading("使用说明", level=2)
            current = "说明"
        elif kind == "正文" and current != "正文":
            doc.add_heading("可替换正文", level=2)
            current = "正文"

        if kind == "formula":
            add_formula(doc, text)
        else:
            add_body(doc, text)

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
