from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


IN_DOC = Path(r"C:/Users/hasee/Desktop/paper_v2_clean_no_wrong_related_comparison.docx")
OUT_DOC = Path(r"C:/Users/hasee/Desktop/paper_v2_reference_methods_comparison.docx")


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=None):
    run.font.name = west
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), east)
    rpr.rFonts.set(qn("w:ascii"), west)
    rpr.rFonts.set(qn("w:hAnsi"), west)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def style_body(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run)


def style_heading(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, east="黑体", size=12, bold=True)


def style_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=10.5)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=9.5, bold=(row_idx == 0))


def find_anchor(doc):
    for p in doc.paragraphs:
        if "由于 AUC 衡量不同阈值下的整体排序能力" in p.text:
            return p
    raise RuntimeError("Could not find section 4.4 anchor.")


def main():
    doc = Document(str(IN_DOC))
    anchor = find_anchor(doc)

    heading = paragraph_after(anchor, "4.4.1 与相关方法的性能比较")
    style_heading(heading)

    intro = paragraph_after(
        heading,
        "为进一步呈现中文有害 Meme 检测任务中的现有方法表现，本文参考文献 [2] 在 ToxiCN MM 测试集上报告的相关结果，"
        "选取文本单模态、图像单模态以及多模态方法进行横向比较。相关方法包括 DeepSeek-V3、RoBERTa、GPT4、CLIP+MKE、VisualBERT COCO、Hate-CLIPper、MOMENTA、PromptHate、Qwen2.5-VL、Debate-based model 和 FG-E2HMD 等。"
    )
    style_body(intro)

    caption = paragraph_after(intro, "表 4-2a 相关方法在 ToxiCN MM 测试集上的性能比较")
    style_caption(caption)

    table = doc.add_table(rows=1, cols=6)
    table._tbl.getparent().remove(table._tbl)
    caption._p.addnext(table._tbl)
    headers = ["类别", "模型", "P", "R", "F1", "F1_harm"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    rows = [
        ["Text-only", "DeepSeek-V3", "74.63", "75.36", "74.99", "58.75"],
        ["Text-only", "RoBERTa", "75.52", "77.54", "76.36", "66.48"],
        ["Text-only", "GPT4", "74.52", "65.59", "68.01", "51.78"],
        ["Image-only", "Image-Region", "65.96", "66.06", "66.01", "52.95"],
        ["Image-only", "ResNet", "66.61", "66.92", "66.76", "53.76"],
        ["Image-only", "ViT", "68.97", "68.61", "68.78", "57.24"],
        ["Multimodal", "GPT4", "74.67", "68.64", "70.11", "55.77"],
        ["Multimodal", "CLIP+MKE", "79.76", "80.79", "80.23", "72.35"],
        ["Multimodal", "VisualBERT COCO", "72.21", "69.36", "70.76", "57.21"],
        ["Multimodal", "Hate-CLIPper", "73.56", "68.52", "70.95", "60.36"],
        ["Multimodal", "MOMENTA", "74.05", "69.88", "71.90", "62.45"],
        ["Multimodal", "PromptHate", "75.83", "72.36", "74.05", "63.15"],
        ["Multimodal", "Qwen2.5-VL", "70.25", "73.36", "71.77", "58.76"],
        ["Multimodal", "Debate-based model", "79.93", "81.16", "79.91", "70.53"],
        ["Multimodal", "FG-E2HMD", "83.53", "83.26", "83.39", "74.43"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)

    analysis = paragraph_after(caption)
    table._tbl.addnext(analysis._p)
    analysis.add_run(
        "从表 4-2a 可以看出，多模态方法整体优于单模态方法，说明图像与文本信息的联合建模对于中文有害 Meme 检测具有重要作用。"
        "其中，CLIP+MKE、Debate-based model 和 FG-E2HMD 等方法取得较高结果，表明显式知识、解释生成和多模态推理能够有效增强模型对隐含有害语义的识别能力。"
        "本文方法同样沿着语义增强的方向展开，但与直接依赖解释文本进行分类不同，SEC-RGCL 将语义解释作为辅助表示信号引入 RGCL 的检索引导对比学习框架，并进一步结合反事实证据学习约束模型对关键判别维度的依赖，从而提升模型在 Acc 和 AUC 指标上的表现。"
    )
    style_body(analysis)

    doc.save(str(OUT_DOC))
    check = Document(str(OUT_DOC))
    print(OUT_DOC)
    print("paragraphs", len(check.paragraphs), "tables", len(check.tables), "images", len(check.inline_shapes))
    print("has reference methods table", any("相关方法在 ToxiCN MM" in p.text for p in check.paragraphs))


if __name__ == "__main__":
    main()
