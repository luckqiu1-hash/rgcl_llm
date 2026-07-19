from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


IN_DOC = Path(r"C:/Users/hasee/Desktop/paper_v2_revised_complete_split_cf_figures_params_source_fixed.docx")
OUT_DOC = Path(r"C:/Users/hasee/Desktop/paper_v2_revised_complete_with_acc_auc_comparison.docx")


def set_run_font(run, east="宋体", west="Times New Roman", size=10.5, bold=None):
    run.font.name = west
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), west)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), west)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def style_body(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run)


def style_heading(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, east="黑体", size=12, bold=True)


def style_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=10.5)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=10, bold=(row_idx == 0))


def find_insert_anchor(doc):
    marker = "两项指标上的一致增益表明"
    for p in doc.paragraphs:
        if marker in p.text:
            return p
    raise RuntimeError("Could not find insertion anchor in section 4.4.")


def main():
    doc = Document(str(IN_DOC))
    anchor = find_insert_anchor(doc)

    heading = paragraph_after(anchor, "4.4.1 与相关方法的性能比较")
    style_heading(heading)

    p1 = paragraph_after(
        heading,
        "为进一步说明本文方法在 ToxiCN MM 数据集上的性能位置，本文参考已有工作 [2] 对该数据集上常用模型的梳理。"
        "该研究将 DeepSeek-V3、RoBERTa、GPT4、CLIP+MKE、VisualBERT COCO、Hate-CLIPper、MOMENTA、PromptHate、Qwen2.5-VL、Debate-based model 和 FG-E2HMD 等方法纳入横向比较，"
        "说明当前中文有害 Meme 检测研究已经从单模态编码逐步转向显式解释、多模态推理和语义增强建模。"
    )
    style_body(p1)

    p2 = paragraph_after(
        p1,
        "在此基础上，本文采用 Acc 和 AUC 作为主要评价指标，对原始 RGCL 基线、辅助文本变体、语义解释变体、反事实变体和完整 SEC-RGCL 进行效果比较。"
        "该对比既能够体现本文方法相对于原始 RGCL 的整体增益，也能够进一步分析语义解释增强和反事实证据学习在分类准确性与排序判别能力上的具体贡献。"
    )
    style_body(p2)

    caption = paragraph_after(p2, "表 4-2a 不同方法的性能比较")
    style_caption(caption)

    table = doc.add_table(rows=1, cols=4)
    table._tbl.getparent().remove(table._tbl)
    caption._p.addnext(table._tbl)

    headers = ["模型", "Acc", "AUC", "说明"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text

    rows = [
        ["RGCL Baseline", "0.8117", "0.8351", "原始图文表示与检索引导对比学习"],
        ["Baseline + Text", "0.8217", "0.8445", "重复引入 OCR 文本作为辅助文本"],
        ["Baseline + Exp", "0.8275", "0.8526", "引入语义解释信息"],
        ["Baseline + CF", "0.8175", "0.8416", "引入反事实证据学习"],
        ["SEC-RGCL", "0.8392", "0.8622", "语义解释增强与反事实证据学习联合建模"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
    style_table(table)

    p3 = paragraph_after(caption)
    table._tbl.addnext(p3._p)
    p3.add_run(
        "从表 4-2a 可以看出，完整 SEC-RGCL 在所有模型变体中取得最优结果。"
        "其中，语义解释分支相较于重复 OCR 文本能够带来更明显的提升，说明模型收益并非来自简单增加文本输入，而是来自对图文组合含义、隐含语义和语境线索的补充建模。"
        "反事实证据学习单独使用时提升幅度相对较小，但与语义解释增强结合后进一步提高整体性能，表明其主要作用在于约束语义增强表示的判别稳定性。"
        "因此，本文方法与已有工作中强调显式解释和语义推理的研究趋势一致，同时在 RGCL 框架内进一步验证了语义增强与证据约束对 Acc 和 AUC 的提升作用。"
    )
    style_body(p3)

    doc.save(str(OUT_DOC))
    check = Document(str(OUT_DOC))
    print(OUT_DOC)
    print("paragraphs", len(check.paragraphs), "tables", len(check.tables), "images", len(check.inline_shapes))
    print("added heading", any("4.4.1 与相关方法" in p.text for p in check.paragraphs))
    print("added table caption", any("表 4-2a" in p.text for p in check.paragraphs))


if __name__ == "__main__":
    main()
