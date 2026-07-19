from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


SRC = Path(r"C:\Users\hasee\Desktop\paper_semantic_refocused_sec_rgcl.docx")
OUT = Path(r"C:\Users\hasee\Desktop\paper_semantic_refocused_with_cf_figure.docx")
FIG = Path(r"D:\pycharm\rgcl_llm\src\cf_evidence_behavior\cf_toxic_only_paper.png")


def set_text(p, text):
    p.clear()
    p.add_run(text)


def insert_paragraph_before(anchor, text="", style=None):
    p = anchor.insert_paragraph_before(text)
    if style:
        try:
            p.style = style
        except Exception:
            pass
    return p


def main():
    doc = Document(str(SRC))

    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("实验在同一组 370 个有害样本上比较完整模型"):
            set_text(
                p,
                "实验在同一组 370 个有害样本上比较完整模型 full 与移除 CF 模块的 no_cf。Mean Drop 和 Median Drop 分别表示置信度下降量的均值与中位数；Drop Rate 表示置信度出现下降的样本比例；Toxic→Non-toxic Flip 表示干预后预测由有害类别翻转为无害类别的样本比例。结果如表 4-5 和图 3 所示。",
            )
            break
    else:
        raise ValueError("Could not find the 4.7 result-introduction paragraph.")

    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("在关键维度被削弱后，full 的 Mean Drop"):
            anchor = p
            break
    if anchor is None:
        raise ValueError("Could not find insertion anchor after Table 4-5.")

    fig_p = insert_paragraph_before(anchor)
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_p.add_run()
    run.add_picture(str(FIG), width=Inches(6.3))

    cap = insert_paragraph_before(
        anchor,
        "图 3. 反事实证据削弱下的有害类别置信度响应。左图展示 full 与 no_cf 在有害样本上的 Drop 分布，其中 Drop = p(toxic | original) − p(toxic | counterfactual)；右图汇总 Mean Drop、Drop Rate 和 Toxic→Non-toxic Flip。更高的下降幅度和翻转比例表明模型对被选中关键证据维度具有更强响应。",
        "Caption",
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
